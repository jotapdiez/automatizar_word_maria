from docx import Document
from docx.shared import Inches

def getOptions(table):
    idx=0
    opciones = ['']*(len(table.row_cells(0))-1)
    for cell in  table.row_cells(0):
        if idx>0 and len(cell.text) > 0:
            opciones[idx - 1] = cell.text.replace('\n', '')
        idx = idx + 1
    return opciones

def getSeccion(table):
    return table.cell(0, 0).text

def getPartes(table):
    idx = 0
    partes =  ['']*(len(table.column_cells(0))-1)
    for cell in table.column_cells(0):
        if idx>0 and len(cell.text) > 0:
            partes[idx - 1] = cell.text.replace('\n', '')
        idx = idx + 1
    return partes

filename = 'Rúbrica de evaluación Dto Naturales EN BLANCO'
document = Document(filename+'.docx')
alumno = input("Alumno: ")
tables = document.tables
for table in tables:
    opciones = getOptions(table)
    seccion = getSeccion(table)
    partes = getPartes(table)

    print('')
    print('')
    print(seccion)
    idx = 1
    respuestas = '0=dejar vacio'
    for opcion in opciones:
        respuestas = respuestas + '|' + str(idx) + '=' + opcion
        idx=idx+1
    respuestas = respuestas + '|' + str(idx) + '= Texto libre'

    idx = 1
    for parte in partes:
        # print(parte)
        valor = input(parte + ' (' + respuestas + "): ")
        if (int(valor) > 0 and int(valor) <= len(opciones)):
            table.cell(idx, int(valor)).text = "X"
        elif (int(valor) > len(opciones)):
            valor = input('Respuesta:')
            table.cell(idx, 1).text = valor
        idx = idx + 1

print('')
print('')
for paragraph in document.paragraphs:
    # format = paragraph.paragraph_format
    # print(paragraph.text+ '- '+ str(paragraph.style.font.bold)+":["+paragraph.style.name+"]")
    if (paragraph.text.strip() == "ALUMNO:"):
        paragraph.add_run(alumno)
        # paragraph.style=paragraph.style
        # paragraph.text = paragraph.text+ '  ' + alumno
        # paragraph.paragraph_format = format
    # print("paragraph:["+paragraph.text.strip()+"]")
    if (paragraph.text.strip() == "EVALUACION FINAL DEL PROCESO:"):
        evaluacion_final = input('EVALUACION FINAL DEL PROCESO:')
        paragraph.add_run(evaluacion_final)
        # paragraph.text = paragraph.text + '  ' + evaluacion_final
        # paragraph.paragraph_format = format
document.save( filename + ' - '+ alumno + '.docx')